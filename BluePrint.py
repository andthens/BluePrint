from flask import Flask, render_template, request, send_file
import os
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def set_font(cell, font_name, font_size, font_color=None, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if font_color:
                run.font.color.rgb = font_color 
            if bold:
                run.font.bold = True  
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)

def set_cell_background(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color) 
    cell_properties.append(cell_shading)

def process_xml(file_path, output_path, in_date=None, user=None, comments=None):
    tree = ET.parse(file_path)
    root = tree.getroot()
    date_format = "%m/%d/%Y"

    doc = Document()
    applet_node = root.find(".//APPLET")
    bc_node = root.find(".//BUSINESS_COMPONENT")

    context_name = None
    if applet_node is not None:
        context_name = applet_node.get("NAME")
        doc.add_heading(f"{context_name}: Controls and Columns", level=1)
        nodes_to_process = [
            {"node": "CONTROL", "attributes": ["NAME", "CAPTION", "HTML_TYPE", "UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "COLUMN", "attributes": ["NAME", "COLUMN_TYPE", "UPDATED", "UPDATED_BY", "COMMENTS"]},
        ]
    elif bc_node is not None:
        context_name = bc_node.get("NAME")
        doc.add_heading(f"{context_name}: Fields", level=1)
        nodes_to_process = [
            {"node": "FIELD", "attributes": ["NAME", "CALCULATED", "CALCULATED_VALUE", "COLUMN", "JOIN", "UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "JOIN", "attributes": ["NAME", "OUTER_JOIN_FLAG", "TABLE","UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "JOIN_SPECIFICATION", "attributes": ["NAME", "DESTINATION_COLUMN","SOURCE_FIELD", "UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "BUSINESS_COMPONENT_USER_PROP", "attributes": ["NAME", "VALUE", "UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "MULTI_VALUE_LINK", "attributes": ["NAME", "DESTINATION_BUSINESS_COMPONENT","DESTINATION_LINK", "UPDATED", "UPDATED_BY"]},
            {"node": "BUSCOMP_SERVER_SCRIPT", "attributes": ["NAME", "SCRIPT", "UPDATED", "UPDATED_BY"]},
        ]
    else:
        return None

    for node_group in nodes_to_process:
        node_name = node_group["node"]
        attributes = node_group["attributes"]

        table_created = False
        table = None

        for element in root.findall(f".//{node_name}"):
            print(node_name)
            updated_date = element.get("UPDATED")
            updated_by = element.get("UPDATED_BY")
            comments_field = element.get("COMMENTS")

            matches = True
            if in_date and updated_date:
                date_only = updated_date.split(" ")[0]
                date_object = datetime.strptime(date_only, date_format)
                matches = matches and date_object > in_date
            if user:
                matches = matches and (updated_by == user)
            if comments:
                matches = matches and (comments in (comments_field or ""))

            if matches:
                if not table_created:
                    doc.add_heading(f"{context_name}: {node_name}", level=1)
                    table = doc.add_table(rows=1, cols=len(attributes))
                    table.style = "Table Grid"

                    hdr_cells = table.rows[0].cells
                    for i, attr in enumerate(attributes):
                        hdr_cells[i].text = attr
                        set_font(hdr_cells[i], "Segoe UI Semilight", 11, font_color=RGBColor(255, 255, 255), bold=True)
                        set_cell_background(hdr_cells[i], "3b6982")
                    table_created = True

                row_cells = table.add_row().cells
                for i, attr in enumerate(attributes):
                    row_cells[i].text = element.get(attr) if element.get(attr) else "N/A"
                    set_font(row_cells[i], "Segoe UI Semilight", 11)

        if table_created and len(table.rows) == 1:
            doc.tables[-1]._element.getparent().remove(doc.tables[-1]._element)

    doc.save(output_path)
    return output_path


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        in_date_str = request.form.get("in_date")
        user = request.form.get("user")
        comments = request.form.get("comments")

        in_date = None
        if in_date_str:
            try:
                in_date = datetime.strptime(in_date_str, "%Y-%m-%d")
            except ValueError:
                return "Invalid date format. Please use the date picker to select a valid date."

        if "file" not in request.files:
            return "No file part"
        file = request.files["file"]
        if file.filename == "":
            return "No selected file"
        if file:
            input_file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(input_file_path)

            output_file_path = os.path.join(OUTPUT_FOLDER, "output.docx")
            processed_file = process_xml(input_file_path, output_file_path, in_date, user, comments)

            if processed_file:
                return send_file(processed_file, as_attachment=True)
            else:
                return "No matching data found in the XML file."
    return render_template("index.html")


if __name__ == "__main__":
    app.run(debug=True)
