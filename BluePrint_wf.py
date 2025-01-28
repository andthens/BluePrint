from flask import Flask, render_template, request, send_file
import os
import xml.etree.ElementTree as ET
from datetime import datetime
from docx import Document
from docx.shared import Pt,RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "outputs"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def add_custom_heading(doc, text, level=1, font_name="Segoe UI Semilight", font_size=10, font_color=None, bold=False):

    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    
    run.font.name = font_name
    run.font.size = Pt(font_size)
    if font_color:
        run.font.color.rgb = font_color
    run.font.bold = bold

    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def set_font(cell, font_name, font_size, font_color=None, bold=False):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            if font_color:
                run.font.color.rgb = font_color
            if bold:
                run.font.bold = True
            r = run._element
            rPr = r.find(qn('w:rPr'))
            if rPr is None:
                rPr = OxmlElement('w:rPr')
                r.insert(0, rPr)
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:ascii'), font_name)
            rFonts.set(qn('w:hAnsi'), font_name)

def set_cell_background(cell, color):
    cell_properties = cell._element.get_or_add_tcPr()
    cell_shading = OxmlElement('w:shd')
    cell_shading.set(qn('w:fill'), color)
    cell_properties.append(cell_shading)

def add_caption_with_seq(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.style = 'Caption' 

    run = paragraph.add_run("Table ")
    run.font.name = "Segoe UI Semilight (Body)"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(59, 105, 130)

    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'SEQ Table \\* ARABIC') 
    fld_run = OxmlElement('w:r')
    fld_text = OxmlElement('w:t')
    fld_text.text = "" 
    fld_run.append(fld_text)
    fld_simple.append(fld_run)
    paragraph._element.append(fld_simple)

    run = paragraph.add_run(f": {text}")
    run.font.name = "Segoe UI Semilight (Body)"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.italic = True
    run.font.color.rgb = RGBColor(59, 105, 130)


def process_xml(file_path, output_path, in_date=None, user=None, comments=None):
    tree = etree.parse(file_path)
    root = tree.getroot()
    date_format = "%m/%d/%Y"

    doc = Document()
    applet_node = root.find(".//APPLET")
    bc_node = root.find(".//BUSINESS_COMPONENT")
    inte_node = root.find(".//INTEGRATION_OBJECT")
    wf_node = root.find(".//WORKFLOW_PROCESS")

    context_name = None
    if applet_node is not None:
        context_name = applet_node.get("NAME")
        nodes_to_process = [
            {"node": "APPLET", "attributes": ["NAME", "TABLE"]},
            {"node": "CONTROL", "attributes": ["NAME", "CAPTION", "HTML_TYPE"]},
            {"node": "COLUMN", "attributes": ["NAME", "COLUMN_TYPE"]},
            {"node": "APPLET_BROWSER_SCRIPT", "attributes": ["NAME", "SCRIPT"]},
            {"node": "APPLET_USER_PROP", "attributes": ["NAME", "VALUE"]},
            {"node": "DRILLDOWN_OBJECT", "attributes": ["NAME", "BUSINESS_COMPONENT", "DESTINATION_FIELD", "SOURCE_FIELD", "HYPERLINK_FIELD"]},
        ]
    elif bc_node is not None:
        context_name = bc_node.get("NAME")
        nodes_to_process = [
            {"node": "BUSINESS_COMPONENT", "attributes": ["NAME", "TABLE"]},
            {"node": "FIELD", "attributes": ["NAME", "CALCULATED", "CALCULATED_VALUE", "COLUMN", "JOIN", "UPDATED", "UPDATED_BY", "COMMENTS"]},
            {"node": "BUSINESS_COMPONENT_USER_PROP", "attributes": ["NAME", "VALUE", "UPDATED", "UPDATED_BY", "COMMENTS"]},
        ]

    elif inte_node is not None:
        context_name = inte_node.get("NAME")
        nodes_to_process = [
            {"node": "INTEGRATION_OBJECT", "attributes": ["NAME", "EXTERNAL_NAME", "XML_TAG"]},
            {"node": "INTEGRATION_COMPONENT", "attributes": ["NAME", "EXTERNAL_NAME", "XML_TAG", "CARDINALITY"]},
            {"node": "INTEGRATION_COMPONENT_FIELD", "attributes": ["NAME", "FIELD_TYPE", "EXTERNAL_NAME", "XML_TAG"]},
            {"node": "INTEGRATION_COMPONENT_KEY", "attributes": ["NAME", "KEY_TYPE"]},
            {"node": "INTEGRATION_COMPONENT_KEY_FIELD", "attributes": ["NAME", "FIELD_NAME"]},
        ]
    elif wf_node is not None:
        context_name = wf_node.get("NAME")
        nodes_to_process = [
            {"node": "WORKFLOW_PROCESS", "attributes": ["NAME", "BUSINESS_OBJECT"]},
            {"node": "WF_STEP", "attributes": ["TYPE","BUSINESS_COMPONENT", "TYPE", "OPERATION"]},
            {"node": "WF_STEP_I_O_ARGUMENT", "attributes": ["NAME","TYPE", "VALUE_SEARCH_SPECIFICATION","BUSINESS_COMPONENT",""]},
            {"node": "INTEGRATION_COMPONENT_KEY", "attributes": ["NAME", "KEY_TYPE"]},
            {"node": "INTEGRATION_COMPONENT_KEY_FIELD", "attributes": ["NAME", "FIELD_NAME"]},
        ]
    else:
        return None

    if context_name:
        add_custom_heading(
            doc,
            text=f"{context_name}",
            level=4,
            font_name="Segoe UI Semilight",
            font_size=10,
            font_color=RGBColor(59, 105, 130),
            bold=False,
        )

    for node_group in nodes_to_process:
        node_name = node_group["node"]
        attributes = node_group["attributes"]

        last_parent = None
        table_created = False
        table = None

        for element in root.findall(f".//{node_name}"):
            updated_date = element.get("UPDATED")
            updated_by = element.get("UPDATED_BY")
            comments_field = element.get("COMMENTS")

            matches = True
            if in_date and updated_date:
                date_only = updated_date.split(" ")[0]
                date_object = datetime.strptime(date_only, date_format)
                matches = matches and date_object > in_date
            if user:
                matches = matches and (updated_by == user)
            if comments:
                matches = matches and (comments in (comments_field or ""))

            parent = element.getparent()

            if matches:
                if not table_created or parent != last_parent:
                    caption_text = f"{context_name} - {node_name} (Parent: {parent.get('NAME', 'N/A') if parent is not None else 'N/A'})"
                    add_caption_with_seq(doc, caption_text)

                    table = doc.add_table(rows=1, cols=len(attributes))
                    table.style = "Table Grid"

                    hdr_cells = table.rows[0].cells
                    for i, attr in enumerate(attributes):
                        hdr_cells[i].text = attr
                        set_font(hdr_cells[i], "Segoe UI Semilight", 11, font_color=RGBColor(255, 255, 255), bold=True)
                        set_cell_background(hdr_cells[i], "3b6982")

                    last_parent = parent
                    table_created = True

                row_cells = table.add_row().cells
                for i, attr in enumerate(attributes):
                    row_cells[i].text = element.get(attr) if element.get(attr) else "N/A"
                    set_font(row_cells[i], "Segoe UI Semilight", 11)

        if table_created and len(table.rows) == 1:
            doc.tables[-1]._element.getparent().remove(doc.tables[-1]._element)

    doc.save(output_path)
    return output_path


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        in_date_str = request.form.get("in_date")
        user = request.form.get("user")
        comments = request.form.get("comments")

        in_date = None
        if in_date_str:
            try:
                in_date = datetime.strptime(in_date_str, "%Y-%m-%d")
            except ValueError:
                return "Invalid date format. Please use the date picker to select a valid date."

        if "file" not in request.files:
            return "No file part"
        file = request.files["file"]
        if file.filename == "":
            return "No selected file"
        if file:
            input_file_path = os.path.join(UPLOAD_FOLDER, file.filename)
            file.save(input_file_path)

            output_file_path = os.path.join(OUTPUT_FOLDER, "output.docx")
            processed_file = process_xml(input_file_path, output_file_path, in_date, user, comments)

            if processed_file:
                return send_file(processed_file, as_attachment=True)
            else:
                return "No matching data found in the XML file."
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
